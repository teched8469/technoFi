from docx import Document
from exceptions import PendingDeprecationWarning

def create_word_document(title, paragraphs, filename):
    doc = Document()
    doc.add_heading(title, 0)

    for para in paragraphs:
        doc.add_paragraph(para)

    doc.save(filename)


title = "My Document"
paragraphs = ["This is the first paragraph.", "This is the second paragraph."]
create_word_document(title, paragraphs, 'document.docx')
